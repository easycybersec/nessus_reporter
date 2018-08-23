import docx 
import pandas as pd 
from docx import Document
from docx.shared import Cm, Inches

doc = Document()

try:
    a = input("Enter name of csv file with format : ")
    d = pd.read_csv(a)
    d = d[~d['Risk'].isin(['None'])]
    d = d.drop_duplicates()
    d = d.groupby(['Name','Risk','Description','Solution','Synopsis','See Also'])['Host'].apply(' '.join).reset_index()
    d = d.rename(columns={'See Also': 'References'})
    d=d.replace(to_replace='Low', value='ZZZ')
    d = d.sort_values('Risk').reset_index()
    d=d.replace(to_replace='ZZZ', value='Low')
    no = d.shape[0]
    d['Host'] = d['Host'].str.split(' ').apply(set).str.join(', ')
    d.to_csv('demo.csv')

    for i in range (no):
        paragraph = doc.add_paragraph(str(1+i)+'.'+d.Name[i])
        table = doc.add_table(rows=6, cols=2)   
        table.style = 'Medium Shading 1 Accent 1'
        for cell in table.columns[1].cells:
            cell.width = Inches(2)
           
        row = table.rows[0]
        row.cells[0].text = 'Host'
        row.cells[1].text = d.Host[i]

        row = table.rows[1]
        row.cells[0].text = 'Synopsis'
        row.cells[1].text = d.Synopsis[i]

        row = table.rows[2]
        row.cells[0].text = 'Risk'
        row.cells[1].text = d.Risk[i]

        row = table.rows[3]
        row.cells[0].text = 'Description'
        row.cells[1].text = str(d.Description[i])

        row = table.rows[4]
        row.cells[0].text = 'Solution'
        row.cells[1].text = d.Solution[i]

        row = table.rows[5]
        row.cells[0].text = 'References'
        row.cells[1].text = d.References[i]
        doc.add_paragraph()

    print("Output saved to file demo.docx, both docx and csv files with same name has been created")  
    doc.save("demo.docx")

except:
    print("Please double check if the entered name is correct.")
